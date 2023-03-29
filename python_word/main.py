import docx

from tkinter import *
from tkinter import ttk


# ====================
# ウィンドウを作成
root = Tk()
# ウィンドウサイズを指定
root.geometry("320x240")
# ウィンドウタイトルを指定
root.title('定期点検')
# ========= フレーム作成
frame1 = ttk.Frame(root, padding=(32))
frame1.grid()
# ラベル 01
label1 = ttk.Label(frame1, text='Cドライブ（DBサーバー）', padding=(5, 2))
label1.grid(row=0, column=0, sticky=E)
# テキストボックス 01
name = StringVar()
name_txt = ttk.Entry(
    frame1,
    textvariable=name,
    width=20)
name_txt.grid(row=0, column=1)
# === OK ボタン
button = ttk.Button(text="OK")
button.place(x=150, y=100)
# ========= ボタン　Click Function =========
def ttx_click():
    w_val = name_txt.get()
    test_l = ttk.Label(text=w_val)
    test_l.place(x=20, y=70)
    print(w_val)
# button に click() function を付加
button["command"] = ttx_click  # 関数と関連付ける

# ウィンドウ表示継続
root.mainloop()

# ==================== word 処理 ====================
# ====== word ファイル　読み込み
doc = docx.Document(
    r"C:\Users\natsume\Desktop\斎場 temp\85-20220819保守サービス報告.docx")

num = 0

# doc.paragraphs[0]
for para in doc.paragraphs:
    num = num + 1
    print(num, para.text)

i = 1
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            print("行:::" + str(i) + ":::" + cell.text)
            i = i + 1
